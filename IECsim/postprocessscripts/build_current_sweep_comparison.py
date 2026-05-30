"""Build current_sweep_comparison.docx for runs 18 (2mA) through 23 (50A) at -100kV.

Sections per spec + extras: overview, x/y/z radial overlays, virtual anode height
vs current (log-log with Gu-Miley threshold and geometric baseline), convergence,
trajdens grids (ions/electrons/combined), epot fields, memory, perveance analysis,
step-count-vs-current cap diagnostic.
"""
import argparse
import os
import re
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# (run_dir, current_A). run18 lives in voltage sweep runs/.
DEFAULT_RUNS = [
    ("voltage sweep runs/run18", 0.002),
    ("current sweep 100kV/run19", 0.010),
    ("current sweep 100kV/run20", 0.100),
    ("current sweep 100kV/run21", 1.000),
    ("current sweep 100kV/run22", 10.000),
    ("current sweep 100kV/run23", 50.000),
]
CATHODE_V = 100000.0  # |V_cathode| in volts (kV*1000)
GEOMETRIC_BASELINE_PCT = 6.864  # vacuum (run18-20) height %; pure geometry
GU_MILEY_K = 0.34  # mA / kV^(3/2)

AXIS_LABELS = {
    "x": "x (side tube)",
    "y": "y (long-tube vertical diameter)",
    "z": "z (long tube)",
}


def fmt_current(I_A):
    if abs(I_A) >= 1.0:
        return f"{I_A:.3g} A"
    if abs(I_A) >= 0.001:
        return f"{I_A * 1000:.3g} mA"
    return f"{I_A * 1e6:.3g} uA"


def perveance_K(I_A, V_kV=100.0):
    """K = I (mA) / V_kV^(3/2). Gu & Miley units."""
    return (I_A * 1000.0) / (V_kV ** 1.5)


def load_epot_error(run):
    p = os.path.join(run, "epot_error.dat")
    if not os.path.exists(p):
        return None
    df = pd.read_csv(p, sep=r"\s+", comment="#", header=None,
                     names=["iteration", "epot_max_error"])
    return df if not df.empty else None


def load_trajectory_steps(run):
    p = os.path.join(run, "trajectory_steps.dat")
    if not os.path.exists(p):
        return None
    df = pd.read_csv(p, sep=r"\s+", comment="#", header=None,
                     names=["iteration", "total_steps"])
    return df if not df.empty else None


def load_radial_final(run, axis="z"):
    p = os.path.join(run, f"potential_radial_{axis}_all.dat")
    if not os.path.exists(p):
        return None
    df = pd.read_csv(p, sep=r"\s+", comment="#", header=None,
                     names=["iteration", "r", "potential"])
    if df.empty:
        return None
    return df[df["iteration"] == df["iteration"].max()].copy()


def load_central_potential_final(run, axis="z"):
    d = load_radial_final(run, axis)
    if d is None:
        return None
    c = d[d["r"].abs() < 1e-10]
    return float(c["potential"].iloc[0]) if len(c) else None


def load_memory_txt(run):
    p = os.path.join(run, "plots", "memory_estimate.txt")
    if not os.path.exists(p):
        return None
    rows = []
    with open(p) as f:
        for ln in f:
            ln = ln.strip()
            if not ln or ln.startswith("#"):
                continue
            parts = ln.split()
            if len(parts) >= 4:
                rows.append(parts)
    return rows or None


def parse_electron_steps_from_log(run):
    """Extract per-iter electron total_steps by parsing fusorsim.log
    (pairs: ion, electron, ion, electron, ...)."""
    p = os.path.join(run, "fusorsim.log")
    if not os.path.exists(p):
        return []
    e_steps = []
    pair = []
    pat = re.compile(r"\s*total\s+steps\s*=\s*(\d+)")
    with open(p) as f:
        for line in f:
            m = pat.match(line)
            if m:
                pair.append(int(m.group(1)))
                if len(pair) == 2:
                    e_steps.append(pair[1])
                    pair = []
    return e_steps


def parse_capped_from_log(run):
    """Sum 'step count limited' counts from log."""
    p = os.path.join(run, "fusorsim.log")
    if not os.path.exists(p):
        return None
    total = 0
    n = 0
    with open(p) as f:
        for line in f:
            m = re.search(r"step count limited\s*=\s*(\d+)", line)
            if m:
                v = int(m.group(1))
                if v > 0:
                    total += v
                    n += 1
    return (total, n)


def make_radial_overlay(runs, out_png, axis="z"):
    fig, ax = plt.subplots(figsize=(11, 7))
    cmap = plt.get_cmap("plasma")
    plotted = 0
    for k, (run, I_A) in enumerate(runs):
        d = load_radial_final(run, axis=axis)
        if d is None:
            continue
        K = perveance_K(I_A)
        color = cmap(k / max(1, len(runs) - 1))
        ax.plot(d["r"], d["potential"], color=color,
                label=f"{fmt_current(I_A)} (K={K:.4g})", linewidth=1.6)
        plotted += 1
    ax.axvline(x=0.005, color="r", linestyle=":", linewidth=2.0,
               label="cathode r=5mm")
    ax.axvline(x=-0.005, color="r", linestyle=":", linewidth=2.0)
    ax.axvline(x=0.01995, color="b", linestyle=":", linewidth=2.0,
               label="anode r=19.95mm")
    ax.axvline(x=-0.01995, color="b", linestyle=":", linewidth=2.0)
    ax.set_xlabel("r (m)")
    ax.set_ylabel("potential (V)")
    ax.set_title(f"Radial Potential Profile - {AXIS_LABELS[axis]} - "
                 f"Current Sweep, -100 kV Deuterium")
    ax.grid(True, alpha=0.3)
    if plotted:
        ax.legend(fontsize=8, loc="best")
    fig.tight_layout()
    fig.savefig(out_png, dpi=150)
    plt.close(fig)


def make_height_vs_K(runs, out_png):
    """KEY PHYSICS PLOT: virtual anode height % vs perveance K (log-log).
    Marks 6.864% geometric baseline, Gu-Miley K=0.34 threshold."""
    K_vals, height_vals = [], []
    for run, I_A in runs:
        V_center = load_central_potential_final(run)
        if V_center is None:
            continue
        height = (CATHODE_V - abs(V_center)) / CATHODE_V * 100.0
        K_vals.append(perveance_K(I_A))
        height_vals.append(height)

    fig, ax = plt.subplots(figsize=(11, 7))
    ax.semilogx(K_vals, height_vals, marker="o", markersize=10,
                color="C3", linewidth=1.8, label="Simulated")
    for K, h, (run, I_A) in zip(K_vals, height_vals, runs):
        ax.annotate(fmt_current(I_A), (K, h), textcoords="offset points",
                    xytext=(10, 10), fontsize=9)
    ax.axhline(y=GEOMETRIC_BASELINE_PCT, color="grey", linestyle="--",
               linewidth=1.2,
               label=f"Geometric baseline ({GEOMETRIC_BASELINE_PCT}%, "
                     f"vacuum field)")
    ax.axvline(x=GU_MILEY_K, color="k", linestyle=":", linewidth=1.2,
               label=f"Gu-Miley threshold K={GU_MILEY_K} (spherical fusor)")
    ax.set_xlabel("Perveance K = I[mA] / V[kV]^(3/2)")
    ax.set_ylabel("Central anode height (% of |V_cathode|)")
    ax.set_title("Virtual Anode Height vs Perveance, -100 kV Deuterium")
    ax.grid(True, alpha=0.3, which="both")
    ax.legend(fontsize=9, loc="best")
    fig.tight_layout()
    fig.savefig(out_png, dpi=150)
    plt.close(fig)


def make_convergence_overlay(runs, out_png):
    fig, ax = plt.subplots(figsize=(11, 7))
    cmap = plt.get_cmap("plasma")
    plotted = 0
    for k, (run, I_A) in enumerate(runs):
        d = load_epot_error(run)
        if d is None:
            continue
        color = cmap(k / max(1, len(runs) - 1))
        ax.plot(d["iteration"], d["epot_max_error"], marker="o",
                color=color, label=fmt_current(I_A), linewidth=1.4)
        plotted += 1
    ax.axhline(y=100.0, color="grey", linestyle="--", linewidth=1.0,
               label="threshold (100 V)")
    ax.set_xlabel("iteration")
    ax.set_ylabel("epot max error (V)")
    ax.set_yscale("symlog", linthresh=0.1)
    ax.set_title("Convergence - Current Sweep, -100 kV Deuterium")
    ax.grid(True, alpha=0.3, which="both")
    if plotted:
        ax.legend(fontsize=9, loc="best")
    fig.tight_layout()
    fig.savefig(out_png, dpi=150)
    plt.close(fig)


def make_step_count_vs_current(runs, out_png):
    """Cap diagnostic: ion + electron step counts (sum across iters) vs current."""
    I_vals, ion_steps, e_steps, caps = [], [], [], []
    for run, I_A in runs:
        ts = load_trajectory_steps(run)
        if ts is None:
            continue
        ion_sum = int(ts["total_steps"].sum())
        e_list = parse_electron_steps_from_log(run)
        e_sum = sum(e_list) if e_list else None
        cap_info = parse_capped_from_log(run)
        cap_sum = cap_info[0] if cap_info else 0
        I_vals.append(I_A)
        ion_steps.append(ion_sum)
        e_steps.append(e_sum if e_sum is not None else 0)
        caps.append(cap_sum)

    fig, ax = plt.subplots(figsize=(11, 7))
    ax.loglog(I_vals, ion_steps, marker="o", linewidth=1.6, label="Ion steps (cumulative)")
    ax.loglog(I_vals, e_steps, marker="s", linewidth=1.6, label="Electron steps (cumulative)")
    ax.loglog(I_vals, caps, marker="^", linewidth=1.6,
              label="'step count limited' hits (cumulative)")
    ax.set_xlabel("Current (A)")
    ax.set_ylabel("count")
    ax.set_title("Trajectory step counts and maxsteps-cap hits vs current")
    ax.grid(True, alpha=0.3, which="both")
    ax.legend(fontsize=9, loc="best")
    fig.tight_layout()
    fig.savefig(out_png, dpi=150)
    plt.close(fig)


def add_overview_table(doc, runs):
    table = doc.add_table(rows=1, cols=9)
    table.style = "Light Grid Accent 1"
    headers = ["Run", "Current", "K = I/V^(3/2)", "Converged",
               "Iters to converge", "Final epot err (V)",
               "Avg ion traj steps", "V(r=0) final (V)",
               "Anode height (%)"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    threshold = 0.001 * CATHODE_V  # 100 V
    for run, I_A in runs:
        err = load_epot_error(run)
        steps = load_trajectory_steps(run)
        V_center = load_central_potential_final(run)
        K = perveance_K(I_A)
        if V_center is not None:
            height = (CATHODE_V - abs(V_center)) / CATHODE_V * 100.0
            vc_str = f"{V_center:.2f}"
            height_str = f"{height:.4f}"
        else:
            vc_str = "-"
            height_str = "-"
        if err is None:
            row = [os.path.basename(run), fmt_current(I_A), f"{K:.4g}",
                   "MISSING", "-", "-",
                   "-" if steps is None else f"{int(steps['total_steps'].mean())}",
                   vc_str, height_str]
        else:
            below = err[err["epot_max_error"] <= threshold]
            if len(below):
                iter_conv = int(below["iteration"].iloc[0])
                converged = "Y"
            else:
                iter_conv = int(err["iteration"].max())
                converged = "N"
            final_err = float(err["epot_max_error"].iloc[-1])
            avg_steps = (f"{int(steps['total_steps'].mean())}"
                         if steps is not None else "-")
            row = [os.path.basename(run), fmt_current(I_A), f"{K:.4g}",
                   converged, str(iter_conv), f"{final_err:.4f}",
                   avg_steps, vc_str, height_str]
        rcells = table.add_row().cells
        for i, v in enumerate(row):
            rcells[i].text = v


def add_image_grid(doc, paths_and_labels, cols=2, width_inches=3.1):
    rows_n = (len(paths_and_labels) + cols - 1) // cols
    table = doc.add_table(rows=rows_n, cols=cols)
    table.autofit = False
    idx = 0
    for r in range(rows_n):
        for c in range(cols):
            if idx >= len(paths_and_labels):
                break
            path, label = paths_and_labels[idx]
            cell = table.rows[r].cells[c]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_obj = p.add_run(label)
            run_obj.bold = True
            run_obj.font.size = Pt(11)
            if os.path.exists(path) and os.path.getsize(path) > 0:
                cell.add_paragraph().add_run().add_picture(
                    path, width=Inches(width_inches))
            else:
                cell.add_paragraph(f"[missing: {path}]")
            idx += 1


def add_memory_summary(doc, runs):
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    headers = ["Run", "Current", "Final iter", "Ion steps",
               "Electron steps", "Total GB"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for run, I_A in runs:
        rows = load_memory_txt(run)
        cells = table.add_row().cells
        cells[0].text = os.path.basename(run)
        cells[1].text = fmt_current(I_A)
        if rows is None:
            for i in range(2, 6):
                cells[i].text = "MISSING"
            continue
        last = rows[-1]
        if len(last) >= 6:
            cells[2].text = last[0]
            cells[3].text = last[1]
            cells[4].text = last[2]
            cells[5].text = last[5]
        else:
            cells[2].text = last[0]
            cells[3].text = last[1]
            cells[4].text = "n/a (legacy)"
            cells[5].text = last[3]


def add_perveance_table(doc, runs):
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    headers = ["Run", "Current", "K", "vs threshold K=0.34",
               "Anode height change vs baseline"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for run, I_A in runs:
        K = perveance_K(I_A)
        ratio = K / GU_MILEY_K
        if ratio < 1:
            cmp_str = f"{1.0/ratio:.4g}x below"
        else:
            cmp_str = f"{ratio:.4g}x above"
        V_center = load_central_potential_final(run)
        if V_center is not None:
            height = (CATHODE_V - abs(V_center)) / CATHODE_V * 100.0
            delta = height - GEOMETRIC_BASELINE_PCT
            delta_str = f"{delta:+.4f}%"
        else:
            delta_str = "-"
        cells = table.add_row().cells
        cells[0].text = os.path.basename(run)
        cells[1].text = fmt_current(I_A)
        cells[2].text = f"{K:.4g}"
        cells[3].text = cmp_str
        cells[4].text = delta_str


def parse_args():
    ap = argparse.ArgumentParser()
    ap.add_argument("--runs", nargs="+", default=None,
                    help="run:current_A pairs, e.g. run19:0.01 run20:0.1 ...")
    ap.add_argument("--out", default="current_sweep_comparison.docx")
    ap.add_argument("--title", default=None)
    return ap.parse_args()


def main():
    args = parse_args()
    if args.runs:
        runs = []
        for entry in args.runs:
            run, cur = entry.rsplit(":", 1)
            runs.append((run, float(cur)))
    else:
        runs = DEFAULT_RUNS

    out_stem = os.path.splitext(args.out)[0]
    doc = Document()
    title = args.title or (
        "Current Sweep Comparison at -100 kV: "
        + ", ".join(fmt_current(c) for _, c in runs)
    )
    doc.add_heading(title, level=0)
    p = doc.add_paragraph(
        "Parameters held constant: -100 kV cathode, deuterium ions and electrons "
        "in matched currents, 800K particles each, h=0.3mm mesh, j<=20 cap, "
        "convergence: 3 consecutive iterations below 100 V epot_max_error."
    )
    p.runs[0].italic = True

    doc.add_heading("1. Overview", level=1)
    add_overview_table(doc, runs)

    doc.add_heading("2. Superimposed Radial Potential Profiles (x, y, z)", level=1)
    for axis in ("x", "y", "z"):
        png = f"{out_stem}_radial_overlay_{axis}.png"
        make_radial_overlay(runs, png, axis=axis)
        if os.path.exists(png):
            doc.add_picture(png, width=Inches(6.5))
            cap = doc.add_paragraph(
                f"{AXIS_LABELS[axis]} axis - final iter per run, "
                f"colored by current (plasma)")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].italic = True
            cap.runs[0].font.size = Pt(9)

    doc.add_heading("3. Virtual Anode Height vs Perveance (KEY PHYSICS)", level=1)
    height_png = f"{out_stem}_height_vs_K.png"
    make_height_vs_K(runs, height_png)
    doc.add_picture(height_png, width=Inches(6.5))
    note = doc.add_paragraph(
        "Anode height = (|V_cathode| - |V(r=0)|) / |V_cathode|. Dashed grey line "
        "= 6.864% geometric baseline (vacuum / no-space-charge limit, from voltage "
        "sweep). Dotted vertical = Gu & Miley (2000) threshold for virtual-anode "
        "formation in a fully spherical fusor; our tee geometry pushes the "
        "effective threshold much higher because ions escape into side tubes "
        "before central space-charge can build up."
    )
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(9)

    doc.add_heading("4. Superimposed Convergence", level=1)
    conv_png = f"{out_stem}_convergence.png"
    make_convergence_overlay(runs, conv_png)
    doc.add_picture(conv_png, width=Inches(6.5))

    doc.add_heading("5. Trajectory Density Comparison (XZ slice)", level=1)
    any_ions = any(os.path.exists(os.path.join(r, "trajdens_ions_xz.png"))
                   for r, _ in runs)
    if any_ions:
        doc.add_heading("5.1 Ions", level=2)
        grid = [(os.path.join(r, "trajdens_ions_xz.png"), fmt_current(I))
                for r, I in runs]
        add_image_grid(doc, grid, cols=2, width_inches=3.0)
        doc.add_heading("5.2 Electrons", level=2)
        grid = [(os.path.join(r, "trajdens_electrons_xz.png"), fmt_current(I))
                for r, I in runs]
        add_image_grid(doc, grid, cols=2, width_inches=3.0)
        doc.add_heading("5.3 Combined (red=ions, blue=electrons, magenta=overlap)",
                        level=2)
        grid = [(os.path.join(r, "plots", "trajdens_combined_xz.png"),
                 fmt_current(I)) for r, I in runs]
        add_image_grid(doc, grid, cols=2, width_inches=3.0)
    else:
        grid = [(os.path.join(r, "trajdens_xz.png"), fmt_current(I))
                for r, I in runs]
        add_image_grid(doc, grid, cols=2, width_inches=3.0)

    doc.add_heading("6. Final Potential Field Comparison (epot_xz)", level=1)
    grid = [(os.path.join(r, "epot_xz.png"), fmt_current(I)) for r, I in runs]
    add_image_grid(doc, grid, cols=2, width_inches=3.0)

    doc.add_heading("7. Memory Summary", level=1)
    add_memory_summary(doc, runs)

    doc.add_heading("8. Perveance Analysis", level=1)
    add_perveance_table(doc, runs)

    doc.add_heading("9. Step-count and Cap Diagnostic vs Current", level=1)
    step_png = f"{out_stem}_step_count_vs_current.png"
    make_step_count_vs_current(runs, step_png)
    doc.add_picture(step_png, width=Inches(6.5))
    note = doc.add_paragraph(
        "Log-log plot of cumulative trajectory steps and IBSimu 'step count "
        "limited' (maxsteps=1000 cap) hits across all iterations of each run. "
        "Cap hits growing super-linearly with current indicates space-charge "
        "trapping is starting to matter."
    )
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(9)

    doc.save(args.out)
    print(f"Wrote {args.out}")


if __name__ == "__main__":
    main()
