"""Build voltage_sweep_comparison.docx aggregating runs 12-18.

Default: runs 12-18 (voltage sweep). Override with CLI args:
    python build_comparison_docx.py --runs run10:-5 run11:-5 --out test.docx
"""
import argparse
import os
import sys
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

DEFAULT_RUNS = [
    ("run12", -10.0),
    ("run13", -20.0),
    ("run14", -30.0),
    ("run15", -40.0),
    ("run16", -50.0),
    ("run17", -75.0),
    ("run18", -100.0),
]
RUNS = DEFAULT_RUNS


def voltage_label(kv):
    return f"{int(kv)} kV"


def load_run_epot_error(run):
    p = os.path.join(run, "epot_error.dat")
    if not os.path.exists(p):
        return None
    df = pd.read_csv(p, sep=r"\s+", comment="#", header=None,
                     names=["iteration", "epot_max_error"])
    if df.empty:
        return None
    return df


def load_run_trajectory_steps(run):
    p = os.path.join(run, "trajectory_steps.dat")
    if not os.path.exists(p):
        return None
    df = pd.read_csv(p, sep=r"\s+", comment="#", header=None,
                     names=["iteration", "total_steps"])
    if df.empty:
        return None
    return df


def load_run_radial_final(run, axis="z"):
    p = os.path.join(run, f"potential_radial_{axis}_all.dat")
    if not os.path.exists(p):
        return None
    df = pd.read_csv(p, sep=r"\s+", comment="#", header=None,
                     names=["iteration", "r", "potential"])
    if df.empty:
        return None
    last_iter = df["iteration"].max()
    return df[df["iteration"] == last_iter].copy()


AXIS_LABELS = {
    "x": "x (side tube)",
    "y": "y (long-tube vertical diameter)",
    "z": "z (long tube)",
}


def parse_central_potential(run):
    """Reads final-iteration V(r=0) per axis from potential_radial_*_all.dat.
    The legacy centre_potential.txt average pools vacuum + converged values
    and is unreliable when the loop needs >1 iter for self-consistency."""
    vals = {}
    for axis in ("x", "y", "z"):
        src = os.path.join(run, f"potential_radial_{axis}_all.dat")
        if not os.path.exists(src):
            continue
        df = pd.read_csv(src, sep=r"\s+", comment="#", header=None,
                         names=["iteration", "r", "potential"])
        if df.empty:
            continue
        last_iter = df["iteration"].max()
        d = df[df["iteration"] == last_iter]
        center = d[d["r"].abs() < 1e-10]
        if len(center):
            vals[axis] = float(center["potential"].iloc[0])
    return vals if vals else None


def central_anode_height_pct(central_v, cathode_v):
    if cathode_v == 0:
        return None
    return (abs(cathode_v) - abs(central_v)) / abs(cathode_v) * 100.0


def load_memory_txt(run):
    p = os.path.join(run, "plots", "memory_estimate.txt")
    if not os.path.exists(p):
        return None
    rows = []
    with open(p) as f:
        for ln in f:
            ln = ln.strip()
            if not ln or ln.startswith("#"):
                continue
            parts = ln.split()
            if len(parts) >= 4:
                rows.append(parts)
    if not rows:
        return None
    return rows


def make_radial_overlay(out_png, axis="z"):
    fig, ax = plt.subplots(figsize=(11, 7))
    cmap = plt.get_cmap("plasma")
    plotted = 0
    for k, (run, kv) in enumerate(RUNS):
        d = load_run_radial_final(run, axis=axis)
        if d is None:
            continue
        color = cmap(k / max(1, len(RUNS) - 1))
        ax.plot(d["r"], d["potential"], color=color,
                label=voltage_label(kv), linewidth=1.6)
        plotted += 1
    ax.axvline(x=0.005, color="r", linestyle=":", linewidth=2.0, label="cathode r=5mm")
    ax.axvline(x=-0.005, color="r", linestyle=":", linewidth=2.0)
    ax.axvline(x=0.01995, color="b", linestyle=":", linewidth=2.0, label="anode r=19.95mm")
    ax.axvline(x=-0.01995, color="b", linestyle=":", linewidth=2.0)
    ax.set_xlabel("r (m)")
    ax.set_ylabel("potential (V)")
    ax.set_title(f"Radial Potential Profile - {AXIS_LABELS[axis]} - Voltage Sweep, 2mA Deuterium")
    ax.grid(True, alpha=0.3)
    if plotted > 0:
        ax.legend(fontsize=9, loc="best")
    fig.tight_layout()
    fig.savefig(out_png, dpi=150)
    plt.close(fig)


def make_convergence_overlay(out_png):
    fig, ax = plt.subplots(figsize=(11, 7))
    cmap = plt.get_cmap("plasma")
    plotted = 0
    for k, (run, kv) in enumerate(RUNS):
        d = load_run_epot_error(run)
        if d is None:
            continue
        color = cmap(k / max(1, len(RUNS) - 1))
        ax.plot(d["iteration"], d["epot_max_error"], marker="o",
                color=color, label=voltage_label(kv), linewidth=1.4)
        plotted += 1
    ax.set_xlabel("iteration")
    ax.set_ylabel("epot max error (V)")
    ax.set_yscale("symlog", linthresh=0.1)
    ax.set_title("Convergence - Voltage Sweep, 2mA Deuterium")
    ax.grid(True, alpha=0.3, which="both")
    if plotted > 0:
        ax.legend(fontsize=9, loc="best")
    fig.tight_layout()
    fig.savefig(out_png, dpi=150)
    plt.close(fig)


def add_overview_table(doc):
    table = doc.add_table(rows=1, cols=8)
    table.style = "Light Grid Accent 1"
    headers = ["Run", "Voltage (kV)", "Converged (Y/N)",
               "Iterations to converge", "Final Epot error (V)",
               "Avg trajectory steps", "Central potential (V)",
               "Central anode height (%)"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for run, kv in RUNS:
        err = load_run_epot_error(run)
        steps = load_run_trajectory_steps(run)
        thresh = 0.001 * abs(kv * 1000.0)
        central = parse_central_potential(run)
        if central is not None:
            avg_central = sum(central.values()) / len(central)
            height_str = f"{central_anode_height_pct(avg_central, kv * 1000.0):.3f}"
            central_str = f"{avg_central:.2f}"
        else:
            avg_central = None
            height_str = "-"
            central_str = "-"
        if err is None:
            row = [run, f"{abs(kv):.0f}", "MISSING", "-", "-",
                   "-" if steps is None else f"{int(steps['total_steps'].mean())}",
                   central_str, height_str]
        else:
            below = err[err["epot_max_error"] <= thresh]
            if len(below) > 0:
                iter_conv = int(below["iteration"].iloc[0])
                converged = "Y"
            else:
                iter_conv = int(err["iteration"].max())
                converged = "N"
            final_err = float(err["epot_max_error"].iloc[-1])
            avg_steps = (f"{int(steps['total_steps'].mean())}"
                         if steps is not None else "-")
            row = [run, f"{abs(kv):.0f}", converged, str(iter_conv),
                   f"{final_err:.4f}", avg_steps, central_str, height_str]
        rcells = table.add_row().cells
        for i, v in enumerate(row):
            rcells[i].text = v


def add_image_grid(doc, paths_and_labels, cols=2, width_inches=3.1):
    """paths_and_labels: list of (image_path, label_text)."""
    rows_n = (len(paths_and_labels) + cols - 1) // cols
    table = doc.add_table(rows=rows_n, cols=cols)
    table.autofit = False
    idx = 0
    for r in range(rows_n):
        for c in range(cols):
            if idx >= len(paths_and_labels):
                break
            path, label = paths_and_labels[idx]
            cell = table.rows[r].cells[c]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_obj = p.add_run(label)
            run_obj.bold = True
            run_obj.font.size = Pt(11)
            if os.path.exists(path) and os.path.getsize(path) > 0:
                cell.add_paragraph().add_run().add_picture(
                    path, width=Inches(width_inches))
            else:
                cell.add_paragraph(f"[missing: {path}]")
            idx += 1


def add_memory_summary(doc):
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    headers = ["Run", "Voltage (kV)", "Final iter",
               "Ion steps", "Electron steps", "Total GB"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for run, kv in RUNS:
        rows = load_memory_txt(run)
        if rows is None:
            cells = table.add_row().cells
            cells[0].text = run
            cells[1].text = f"{abs(kv):.0f}"
            for i in range(2, 6):
                cells[i].text = "MISSING"
            continue
        last = rows[-1]
        cells = table.add_row().cells
        cells[0].text = run
        cells[1].text = f"{abs(kv):.0f}"
        if len(last) >= 6:
            # New 6-col format: iter ion_steps e_steps traj_MB overhead_MB total_GB
            cells[2].text = last[0]
            cells[3].text = last[1]
            cells[4].text = last[2]
            cells[5].text = last[5]
        else:
            # Legacy 4-col format: iter total_steps total_MB total_GB
            cells[2].text = last[0]
            cells[3].text = last[1]
            cells[4].text = "n/a (legacy)"
            cells[5].text = last[3]


def parse_args():
    ap = argparse.ArgumentParser()
    ap.add_argument("--runs", nargs="+", default=None,
                    help="List of run:kV pairs, e.g. run12:-10 run13:-20")
    ap.add_argument("--out", default="voltage_sweep_comparison.docx",
                    help="Output docx path")
    ap.add_argument("--title", default=None,
                    help="Override heading title")
    return ap.parse_args()


def main():
    global RUNS
    args = parse_args()
    if args.runs:
        RUNS = []
        for entry in args.runs:
            run, kv = entry.split(":")
            RUNS.append((run, float(kv)))
    title = args.title or f"Voltage Sweep Comparison: {', '.join(r for r, _ in RUNS)}"
    doc = Document()
    doc.add_heading(title, level=0)
    p = doc.add_paragraph(
        "Parameters held constant: 2mA ion + 2mA electron current, "
        "800K particles each, deuterium ions, h=0.3mm mesh, 10-iter cap, "
        "convergence threshold = 0.1% of |cathodepot|."
    )
    p.runs[0].italic = True

    doc.add_heading("1. Overview", level=1)
    add_overview_table(doc)

    out_stem = os.path.splitext(args.out)[0]
    doc.add_heading("2. Superimposed Radial Potential Profiles (x, y, z)", level=1)
    for axis in ("x", "y", "z"):
        radial_png = f"{out_stem}_radial_overlay_{axis}.png"
        make_radial_overlay(radial_png, axis=axis)
        if os.path.exists(radial_png):
            doc.add_picture(radial_png, width=Inches(6.5))
            cap = doc.add_paragraph(f"{AXIS_LABELS[axis]} axis - final iteration per run")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].italic = True
            cap.runs[0].font.size = Pt(9)

    doc.add_heading("3. Superimposed Convergence", level=1)
    conv_png = f"{out_stem}_convergence_overlay.png"
    make_convergence_overlay(conv_png)
    if os.path.exists(conv_png):
        doc.add_picture(conv_png, width=Inches(6.5))

    doc.add_heading("4. Trajectory Density Comparison (XZ slice)", level=1)
    any_ion_split = any(
        os.path.exists(os.path.join(run, "trajdens_ions_xz.png"))
        for run, _ in RUNS
    )
    if any_ion_split:
        doc.add_heading("4.1 Ions", level=2)
        grid = [(os.path.join(run, "trajdens_ions_xz.png"), voltage_label(kv))
                for run, kv in RUNS]
        add_image_grid(doc, grid, cols=2, width_inches=3.0)

        doc.add_heading("4.2 Electrons", level=2)
        grid = [(os.path.join(run, "trajdens_electrons_xz.png"), voltage_label(kv))
                for run, kv in RUNS]
        add_image_grid(doc, grid, cols=2, width_inches=3.0)

        doc.add_heading("4.3 Combined (red=ions, blue=electrons, magenta=overlap)",
                        level=2)
        grid = [(os.path.join(run, "plots", "trajdens_combined_xz.png"),
                 voltage_label(kv))
                for run, kv in RUNS]
        add_image_grid(doc, grid, cols=2, width_inches=3.0)
    else:
        grid = [(os.path.join(run, "trajdens_xz.png"), voltage_label(kv))
                for run, kv in RUNS]
        add_image_grid(doc, grid, cols=2, width_inches=3.0)

    doc.add_heading("5. Final Potential Field Comparison (epot_xz)", level=1)
    grid2 = [(os.path.join(run, "epot_xz.png"), voltage_label(kv))
             for run, kv in RUNS]
    add_image_grid(doc, grid2, cols=2, width_inches=3.0)

    doc.add_heading("6. Memory Summary", level=1)
    add_memory_summary(doc)

    doc.save(args.out)
    print(f"Wrote {args.out}")


if __name__ == "__main__":
    main()
