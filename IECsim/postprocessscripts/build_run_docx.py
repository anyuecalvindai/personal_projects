"""Build per-run summary docx for one voltage-sweep run.

Usage: python build_run_docx.py runN
Reads files from runN/ and writes runN/runN_summary.docx.
"""
import os
import sys
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _cpp_path_for(run_dir):
    base = os.path.basename(run_dir.rstrip("/"))
    cpp_path = os.path.join(run_dir, f"fusorsim{base}.cpp")
    if not os.path.exists(cpp_path):
        cpp_path = os.path.join(run_dir, f"fusorsimrun{base}.cpp")
    return cpp_path if os.path.exists(cpp_path) else None


def voltage_kv_from_run(run_dir):
    cpp_path = _cpp_path_for(run_dir)
    if cpp_path:
        with open(cpp_path) as f:
            for line in f:
                if line.strip().startswith("double cathodepot"):
                    val = line.split("=")[1].split(";")[0].strip()
                    try:
                        return float(val) / 1000.0
                    except ValueError:
                        pass
    return None


def beam_current_A_from_run(run_dir):
    cpp_path = _cpp_path_for(run_dir)
    if not cpp_path:
        return None
    with open(cpp_path) as f:
        for line in f:
            s = line.strip()
            if s.startswith("const double beam_current ") or s.startswith("const double beam_current="):
                val = line.split("=")[1].split(";")[0].strip()
                try:
                    return float(val)
                except ValueError:
                    pass
    return None


def format_current(I_A):
    if I_A is None:
        return "unknown"
    if abs(I_A) >= 1.0:
        return f"{I_A:.3g} A"
    if abs(I_A) >= 0.001:
        return f"{I_A * 1000:.3g} mA"
    return f"{I_A * 1e6:.3g} uA"


AXIS_LABELS = {
    "x": "x (side tube)",
    "y": "y (long-tube vertical diameter)",
    "z": "z (long tube)",
}


def make_radial_overlay(run_dir, axis, out_png):
    src = os.path.join(run_dir, f"potential_radial_{axis}_all.dat")
    if not os.path.exists(src):
        return False
    df = pd.read_csv(src, sep=r"\s+", comment="#", header=None,
                     names=["iteration", "r", "potential"])
    fig, ax = plt.subplots(figsize=(10, 6))
    iters = sorted(df["iteration"].unique())
    cmap = plt.get_cmap("viridis")
    for k, it in enumerate(iters):
        d = df[df["iteration"] == it]
        ax.plot(d["r"], d["potential"], color=cmap(k / max(1, len(iters) - 1)),
                label=f"iter {int(it)}", linewidth=1.4)
    ax.set_xlabel("r (m)")
    ax.set_ylabel("potential (V)")
    ax.set_title(f"Radial potential profile ({AXIS_LABELS[axis]}) - all iterations")
    ax.grid(True, alpha=0.3)
    ax.axvline(x=0, color="k", linestyle="--", linewidth=0.5)
    ax.axvline(x=0.005, color="r", linestyle=":", linewidth=0.8, label="cathode r=5mm")
    ax.axvline(x=-0.005, color="r", linestyle=":", linewidth=0.8)
    ax.axvline(x=0.01995, color="b", linestyle=":", linewidth=0.8, label="anode r=19.95mm")
    ax.axvline(x=-0.01995, color="b", linestyle=":", linewidth=0.8)
    ax.legend(fontsize=7, ncol=2)
    fig.tight_layout()
    fig.savefig(out_png, dpi=150)
    plt.close(fig)
    return True


def make_iter1_vs_final_overlay(run_dir, axis, out_png):
    """Vacuum (iter 1) vs self-consistent (last iter) radial profile - the gap = space-charge effect."""
    src = os.path.join(run_dir, f"potential_radial_{axis}_all.dat")
    if not os.path.exists(src):
        return False
    df = pd.read_csv(src, sep=r"\s+", comment="#", header=None,
                     names=["iteration", "r", "potential"])
    iters = sorted(df["iteration"].unique())
    if len(iters) < 1:
        return False
    first = df[df["iteration"] == iters[0]]
    last = df[df["iteration"] == iters[-1]]
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.plot(first["r"], first["potential"], color="grey", linewidth=1.4,
            linestyle="--", label=f"iter {int(iters[0])} (vacuum field)")
    ax.plot(last["r"], last["potential"], color="C3", linewidth=1.8,
            label=f"iter {int(iters[-1])} (self-consistent)")
    ax.set_xlabel("r (m)")
    ax.set_ylabel("potential (V)")
    ax.set_title(f"Iter-1 vs final radial profile ({AXIS_LABELS[axis]})")
    ax.grid(True, alpha=0.3)
    ax.axvline(x=0.005, color="r", linestyle=":", linewidth=0.8, label="cathode r=5mm")
    ax.axvline(x=-0.005, color="r", linestyle=":", linewidth=0.8)
    ax.axvline(x=0.01995, color="b", linestyle=":", linewidth=0.8, label="anode r=19.95mm")
    ax.axvline(x=-0.01995, color="b", linestyle=":", linewidth=0.8)
    ax.legend(fontsize=8)
    fig.tight_layout()
    fig.savefig(out_png, dpi=150)
    plt.close(fig)
    return True


def find_extrema(run_dir, axis="z", min_prominence_V=None):
    """Find significant extrema in the final radial profile.
    Uses scipy.signal.find_peaks with prominence filter to suppress grid noise.
    min_prominence_V: required peak/trough prominence in V. None => no filter.
    Returns list of (r, V, kind) tuples sorted by r."""
    from scipy.signal import find_peaks
    src = os.path.join(run_dir, f"potential_radial_{axis}_all.dat")
    if not os.path.exists(src):
        return None
    df = pd.read_csv(src, sep=r"\s+", comment="#", header=None,
                     names=["iteration", "r", "potential"])
    last_iter = df["iteration"].max()
    d = df[df["iteration"] == last_iter].sort_values("r").reset_index(drop=True)
    r = d["r"].to_numpy()
    V = d["potential"].to_numpy()
    prom = min_prominence_V if min_prominence_V is not None else None
    max_idx, _ = find_peaks(V, prominence=prom)
    min_idx, _ = find_peaks(-V, prominence=prom)
    extrema = [(r[i], V[i], "max") for i in max_idx] + \
              [(r[i], V[i], "min") for i in min_idx]
    extrema.sort(key=lambda t: t[0])
    return extrema


def make_convergence_inset(run_dir, out_png):
    """Per-iter epot_max_error trajectory; flags oscillation."""
    src = os.path.join(run_dir, "epot_error.dat")
    if not os.path.exists(src):
        return False
    df = pd.read_csv(src, sep=r"\s+", comment="#", header=None,
                     names=["iteration", "epot_max_error"])
    if df.empty:
        return False
    fig, ax = plt.subplots(figsize=(8, 4))
    ax.plot(df["iteration"], df["epot_max_error"], marker="o", linewidth=1.5,
            color="C2")
    ax.set_xlabel("iteration")
    ax.set_ylabel("Epot max error (V)")
    # Detect oscillation: any increase between consecutive iterations
    errs = df["epot_max_error"].to_numpy()
    oscillating = any(errs[i] > errs[i - 1] for i in range(1, len(errs)))
    title = "Convergence per iteration"
    if oscillating:
        title += " - OSCILLATING (error increased between iters)"
    ax.set_title(title)
    ax.grid(True, alpha=0.3)
    if (errs > 0).any() and (errs >= 0).all():
        ax.set_yscale("symlog", linthresh=0.1)
    fig.tight_layout()
    fig.savefig(out_png, dpi=150)
    plt.close(fig)
    return True


def parse_central_potential(run_dir):
    """Reads final-iteration V(r=0) for each axis directly from the
    potential_radial_{axis}_all.dat files. Returns dict {'x': V, 'y': V, 'z': V}.
    The legacy centre_potential.txt average is unreliable under the new
    convergence criterion because it pools the vacuum (iter 1) value with
    the converged (later iter) values."""
    vals = {}
    for axis in ("x", "y", "z"):
        src = os.path.join(run_dir, f"potential_radial_{axis}_all.dat")
        if not os.path.exists(src):
            continue
        df = pd.read_csv(src, sep=r"\s+", comment="#", header=None,
                         names=["iteration", "r", "potential"])
        if df.empty:
            continue
        last_iter = df["iteration"].max()
        d = df[df["iteration"] == last_iter]
        center = d[d["r"].abs() < 1e-10]
        if len(center):
            vals[axis] = float(center["potential"].iloc[0])
    return vals if vals else None


def central_anode_height_pct(central_v, cathode_v):
    """Height % = (|cathodepot| - |central|) / |cathodepot| * 100.
    Positive means central potential is LESS negative than cathode (virtual-anode lift)."""
    if cathode_v == 0:
        return None
    return (abs(cathode_v) - abs(central_v)) / abs(cathode_v) * 100.0


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_image(doc, path, width_inches=6.0, caption=None):
    if os.path.exists(path) and os.path.getsize(path) > 0:
        doc.add_picture(path, width=Inches(width_inches))
        if caption:
            p = doc.paragraphs[-1]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].italic = True
            cap.runs[0].font.size = Pt(9)
        return True
    p = doc.add_paragraph(f"[missing: {path}]")
    p.runs[0].italic = True
    return False


def add_memory_table(doc, mem_path):
    if not os.path.exists(mem_path):
        doc.add_paragraph(f"[missing: {mem_path}]").runs[0].italic = True
        return
    header_lines = []
    rows = []
    with open(mem_path) as f:
        for ln in f:
            ln = ln.rstrip()
            if not ln.strip():
                continue
            if ln.startswith("#"):
                cleaned = ln.lstrip("#").strip()
                if cleaned and not cleaned.startswith("iter"):
                    header_lines.append(cleaned)
                continue
            parts = ln.split()
            rows.append(parts)
    for h in header_lines:
        p = doc.add_paragraph(h)
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(9)
    if not rows:
        doc.add_paragraph("[memory_estimate.txt empty]")
        return
    ncols = len(rows[0])
    if ncols >= 6:
        header_cells = ["iteration", "ion_steps", "e_steps",
                        "traj_MB", "overhead_MB", "total_GB"]
    elif ncols == 4:
        header_cells = ["iteration", "total_steps", "total_MB", "total_GB"]
    else:
        header_cells = [f"col{i}" for i in range(ncols)]
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(header_cells[:ncols]):
        table.rows[0].cells[i].text = h
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row[:ncols]):
            table.rows[ri].cells[ci].text = val


def main(run_dir):
    if not run_dir.endswith("/"):
        run_dir = run_dir + "/"
    run_base = os.path.basename(run_dir.rstrip("/"))
    run_suffix = run_base[3:] if run_base.startswith("run") else run_base

    voltage_kv = voltage_kv_from_run(run_dir)
    voltage_v = voltage_kv * 1000.0 if voltage_kv is not None else None
    voltage_str = f"{voltage_kv:.0f} kV" if voltage_kv is not None else "unknown"
    beam_A = beam_current_A_from_run(run_dir)
    current_str = format_current(beam_A)

    os.makedirs(os.path.join(run_dir, "plots"), exist_ok=True)
    overlay_pngs = {}
    iter1_vs_final_pngs = {}
    for axis in ("x", "y", "z"):
        out = os.path.join(run_dir, "plots", f"radial_potential_overlay_{axis}.png")
        if make_radial_overlay(run_dir, axis, out):
            overlay_pngs[axis] = out
        iv_out = os.path.join(run_dir, "plots", f"iter1_vs_final_{axis}.png")
        if make_iter1_vs_final_overlay(run_dir, axis, iv_out):
            iter1_vs_final_pngs[axis] = iv_out
    conv_inset_png = os.path.join(run_dir, "plots", "convergence_inset.png")
    if not make_convergence_inset(run_dir, conv_inset_png):
        conv_inset_png = None

    doc = Document()
    doc.add_heading(
        f"Run {run_suffix} - {current_str} / {voltage_str}, 800K particles, Deuterium",
        level=0,
    )

    add_heading(doc, "Central Anode Height", level=1)
    central = parse_central_potential(run_dir)
    if central and voltage_v is not None:
        avg_central = sum(central.values()) / len(central)
        height_pct = central_anode_height_pct(avg_central, voltage_v)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Light Grid Accent 1"
        rows_data = [
            ("Cathode voltage", f"{voltage_v:.1f} V ({voltage_str})"),
            ("Central potential (avg of x/y/z probes)", f"{avg_central:.2f} V"),
            ("Central potential (x axis)", f"{central.get('x', 0):.2f} V"),
            ("Central potential (y axis)", f"{central.get('y', 0):.2f} V"),
            ("Central potential (z axis)", f"{central.get('z', 0):.2f} V"),
            ("Central anode height", f"{height_pct:.3f} % of |V_cathode|"),
        ]
        tbl.rows[0].cells[0].text = rows_data[0][0]
        tbl.rows[0].cells[1].text = rows_data[0][1]
        for label, value in rows_data[1:]:
            cells = tbl.add_row().cells
            cells[0].text = label
            cells[1].text = value
        note = doc.add_paragraph(
            "Height % = (|V_cathode| - |V_centre|) / |V_cathode|. "
            "Positive value = central potential is LESS negative than the cathode "
            "(space-charge / geometric 'virtual anode' lift)."
        )
        note.runs[0].italic = True
        note.runs[0].font.size = Pt(9)
    else:
        doc.add_paragraph("[centre_potential.txt or cathodepot not parseable]")

    # Extrema scan - find all radial maxima/minima (for double-well detection)
    add_heading(doc, "Radial Profile Extrema (z-axis, final iter)", level=1)
    extrema_threshold = abs(voltage_v) * 0.01 if voltage_v is not None else None
    extrema = find_extrema(run_dir, axis="z",
                           min_prominence_V=extrema_threshold)
    if extrema_threshold is not None:
        thresh_note = doc.add_paragraph(
            f"Prominence filter: only extrema with prominence > "
            f"{extrema_threshold:.1f} V (1% of |V_cathode|) are reported "
            f"to suppress mesh-resolution noise."
        )
        thresh_note.runs[0].italic = True
        thresh_note.runs[0].font.size = Pt(9)
    if extrema is None:
        doc.add_paragraph("[potential_radial_z_all.dat missing]")
    elif not extrema:
        p = doc.add_paragraph(
            "No significant interior extrema in final radial profile - "
            "monotonic well (no virtual anode or virtual cathode formation "
            "above the 1% prominence threshold)."
        )
        p.runs[0].italic = True
    else:
        etbl = doc.add_table(rows=1, cols=4)
        etbl.style = "Light Grid Accent 1"
        for i, h in enumerate(["#", "r (m)", "V (V)", "kind"]):
            etbl.rows[0].cells[i].text = h
        for i, (r, V, kind) in enumerate(extrema, start=1):
            cells = etbl.add_row().cells
            cells[0].text = str(i)
            cells[1].text = f"{r:+.4f}"
            cells[2].text = f"{V:.2f}"
            cells[3].text = kind
        note = doc.add_paragraph(
            "An interior local MAX inside the cathode (|r| < 5 mm) is a "
            "virtual-anode signature; an interior local MIN above the "
            "cathode floor is a virtual-cathode (double-well) signature."
        )
        note.runs[0].italic = True
        note.runs[0].font.size = Pt(9)

    add_heading(doc, "Iter-1 (vacuum) vs Final (self-consistent) profile", level=1)
    for axis in ("x", "y", "z"):
        if axis in iter1_vs_final_pngs:
            add_image(doc, iter1_vs_final_pngs[axis], width_inches=6.5,
                      caption=f"{AXIS_LABELS[axis]} - gap between dashed and solid = space-charge contribution")
        else:
            doc.add_paragraph(f"[missing potential_radial_{axis}_all.dat]")

    add_heading(doc, "Radial Potential Profiles", level=1)
    for axis in ("x", "y", "z"):
        if axis in overlay_pngs:
            add_image(doc, overlay_pngs[axis], width_inches=6.5,
                      caption=f"Radial potential along {AXIS_LABELS[axis]}, all iterations overlaid")
        else:
            doc.add_paragraph(f"[missing potential_radial_{axis}_all.dat]")

    add_heading(doc, "Final Potential Field", level=1)
    add_image(doc, os.path.join(run_dir, "epot_xz.png"), caption="epot XZ slice")
    add_image(doc, os.path.join(run_dir, "epot_xy.png"), caption="epot XY slice")
    add_image(doc, os.path.join(run_dir, "epot_yz_x0.png"), caption="epot YZ slice at x=0")

    add_heading(doc, "Trajectory Density", level=1)
    trajdens_views = ("xz", "xy", "yz_x0")
    has_ion_split = any(
        os.path.exists(os.path.join(run_dir, f"trajdens_ions_{v}.png"))
        for v in trajdens_views
    )
    if has_ion_split:
        add_heading(doc, "Ions", level=2)
        for v in trajdens_views:
            add_image(doc, os.path.join(run_dir, f"trajdens_ions_{v}.png"),
                      caption=f"Ion trajectory density - {v.upper()} slice")
        add_heading(doc, "Electrons", level=2)
        for v in trajdens_views:
            add_image(doc, os.path.join(run_dir, f"trajdens_electrons_{v}.png"),
                      caption=f"Electron trajectory density - {v.upper()} slice")
        add_heading(doc, "Combined (ions=red, electrons=blue, overlap=magenta)", level=2)
        for v in trajdens_views:
            add_image(doc, os.path.join(run_dir, "plots", f"trajdens_combined_{v}.png"),
                      caption=f"Combined trajectory density - {v.upper()} slice")
    else:
        # Legacy single-pdb trajdens output
        for v in ("xz", "xy"):
            add_image(doc, os.path.join(run_dir, f"trajdens_{v}.png"),
                      caption=f"Trajectory density {v.upper()} slice")

    add_heading(doc, "Memory Estimate", level=1)
    add_memory_table(doc, os.path.join(run_dir, "plots", "memory_estimate.txt"))

    add_heading(doc, "Convergence", level=1)
    if conv_inset_png:
        add_image(doc, conv_inset_png, width_inches=6.5,
                  caption="Epot max error per iteration (oscillation flag in title if error grew between iters)")
    add_image(doc, os.path.join(run_dir, "plots", "epot_convergence.png"),
              caption="Epot max error vs iteration (postprocess plot)")

    out_path = os.path.join(run_dir, f"{run_base}_summary.docx")
    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("usage: build_run_docx.py <run_dir>", file=sys.stderr)
        sys.exit(1)
    main(sys.argv[1])
